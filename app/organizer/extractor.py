"""
Pulls whatever signal exists about "what this document is about" out of a .docx
file: its title metadata (if the author ever set one), its headings, and its
body text. No LLM, no network call -- just reading what's already in the file.

Scoped to .docx for v1 (that's the actual pain point described: "I make a Word
doc, save it as doc1, can never find it again"). PDF/xlsx/pptx extraction would
follow the same shape -- a per-format extractor returning the same
DocumentContent structure -- but aren't built here to avoid speculative scope.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    title_metadata: str = ""       # docx core properties "title" field, if set
    headings: list = field(default_factory=list)   # paragraphs styled as Heading 1/2/...
    body_text: str = ""            # all paragraph text, joined -- what the namer's
                                    # keyword extraction actually runs over


class ExtractionError(Exception):
    pass


def extract_docx(path: Path) -> DocumentContent:
    """Raises ExtractionError on a corrupt/unreadable file rather than returning
    something misleading -- a file that can't be read shouldn't get renamed
    based on a guess."""
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ExtractionError(f"Could not open {path.name} as a .docx file: {e}")

    title_metadata = (doc.core_properties.title or "").strip()

    headings = []
    body_parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        if style_name.lower().startswith("heading") or style_name.lower() == "title":
            headings.append(text)
        body_parts.append(text)

    # doc.paragraphs only covers free-standing body paragraphs -- text inside
    # tables lives in a completely separate doc.tables collection and is
    # invisible to the loop above. This matters a lot in practice: bulleted
    # lists, checklists, and agendas are frequently laid out as a table for
    # alignment, so a table-only document (no free-standing paragraphs) used
    # to extract as effectively empty, silently falling all the way through
    # generate_name()'s fallbacks to a timestamped placeholder even though
    # the file clearly has real content.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    body_parts.append(cell_text)

    return DocumentContent(
        title_metadata=title_metadata,
        headings=headings,
        body_text="\n".join(body_parts),
    )
