from pathlib import Path

from docx import Document

from app.organizer.extractor import extract_docx, ExtractionError
from app.organizer.namer import generate_name, _dedupe_words_preserve_order
from app.organizer.renamer import is_generically_named, safe_rename


def _make_docx(path: Path, title: str = "", heading: str = "", paragraphs=None):
    doc = Document()
    if title:
        doc.core_properties.title = title
    if heading:
        doc.add_heading(heading, level=1)
    for p in paragraphs or []:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def test_extractor_reads_title_metadata_and_headings(tmp_path):
    path = _make_docx(
        tmp_path / "a.docx",
        title="Q3 Marketing Budget Proposal",
        heading="Executive Summary",
        paragraphs=["Some body text."],
    )
    content = extract_docx(path)
    assert content.title_metadata == "Q3 Marketing Budget Proposal"
    assert content.headings == ["Executive Summary"]
    assert "Some body text." in content.body_text


def test_extractor_reads_text_from_tables(tmp_path):
    """Regression test for a real bug: python-docx keeps table content in a
    separate doc.tables collection, invisible to doc.paragraphs. A document
    whose content lives entirely in a table (common for bulleted/checklist
    layouts) used to extract as empty -- no title, no headings, no body text
    -- and generate_name() would silently fall all the way through to a
    timestamped placeholder even though the file clearly has real content."""
    path = tmp_path / "table_only.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Bulletin Board"
    table.cell(0, 1).text = "Meeting Notes"
    table.cell(1, 0).text = "Q3 budget review"
    table.cell(1, 1).text = "approved unanimously"
    doc.save(str(path))

    content = extract_docx(path)
    assert content.body_text.strip() != ""
    assert "Bulletin Board" in content.body_text
    assert "Q3 budget review" in content.body_text


def test_extractor_raises_on_non_docx_file(tmp_path):
    path = tmp_path / "not_really_a_docx.docx"
    path.write_text("this is plain text, not a zip/docx", encoding="utf-8")
    try:
        extract_docx(path)
        assert False, "expected ExtractionError"
    except ExtractionError:
        pass


def test_namer_prefers_title_metadata_over_everything_else():
    from app.organizer.extractor import DocumentContent
    content = DocumentContent(
        title_metadata="Q3 Marketing Budget Proposal",
        headings=["Some Heading"],
        body_text="unrelated body text about something else entirely",
    )
    assert generate_name(content) == "Q3 Marketing Budget Proposal"


def test_namer_falls_back_to_heading_when_no_title():
    from app.organizer.extractor import DocumentContent
    content = DocumentContent(headings=["Employee Onboarding Checklist"], body_text="Step one.")
    assert generate_name(content) == "Employee Onboarding Checklist"


def test_namer_keyword_fallback_has_no_duplicate_words(monkeypatch):
    """Regression guard: overlapping yake bigrams ('annual budget' + 'budget
    review') used to concatenate into 'Annual Budget Budget Review Review
    Process' -- visibly broken. Word-level dedup should collapse that.

    LLM naming is disabled here (monkeypatched to None) so this test
    deterministically exercises the keyword-extraction tier regardless of
    whether Ollama happens to be running on whatever machine runs the tests
    -- that tier is covered separately in test_namer_llm.py."""
    import app.organizer.namer as namer_module
    monkeypatch.setattr(namer_module, "_generate_name_llm", lambda *a, **kw: None)

    from app.organizer.extractor import DocumentContent
    content = DocumentContent(
        body_text=(
            "Following up on our conversation about the annual budget review process. "
            "The finance team needs the quarterly expense reports submitted by Friday "
            "for the annual budget review meeting next week. Please send your department "
            "expense reports as soon as possible so we can finalize the annual budget review."
        )
    )
    name = generate_name(content)
    words = [w.lower() for w in name.split()]
    assert len(words) == len(set(words)), f"name has duplicate words: {name!r}"


def test_namer_dedupe_word_helper_preserves_first_occurrence_order():
    result = _dedupe_words_preserve_order(["Annual Budget", "Budget Review", "Review Process"], max_words=10)
    assert result == ["Annual", "Budget", "Review", "Process"]


def test_namer_falls_back_to_timestamp_when_nothing_extractable():
    from app.organizer.extractor import DocumentContent
    content = DocumentContent()  # empty title, headings, body
    name = generate_name(content)
    assert name.startswith("Untitled Document")


def test_is_generically_named_matches_common_default_names():
    for name in ["doc1", "Document1", "Document (2)", "New Microsoft Word Document", "Untitled", "Untitled (3)"]:
        assert is_generically_named(name), f"expected {name!r} to be recognized as generic"


def test_is_generically_named_rejects_real_titles():
    for name in ["Family Vacation Itinerary", "Q3 Marketing Budget Proposal", "Resume - Jane Doe"]:
        assert not is_generically_named(name), f"expected {name!r} to NOT be recognized as generic"


def test_safe_rename_skips_non_generic_names_by_default(tmp_path):
    path = _make_docx(tmp_path / "Family Vacation Itinerary.docx", paragraphs=["content"])
    log_path = tmp_path / "log.json"
    result = safe_rename(path, "Some New Name", log_path)
    assert result.status == "skipped_not_generic"
    assert path.exists()  # untouched


def test_safe_rename_actually_renames_generic_files(tmp_path):
    path = _make_docx(tmp_path / "doc1.docx", paragraphs=["content"])
    log_path = tmp_path / "log.json"
    result = safe_rename(path, "Q3 Marketing Budget Proposal", log_path)
    assert result.status == "renamed"
    assert result.new_path.name == "Q3 Marketing Budget Proposal.docx"
    assert result.new_path.exists()
    assert not path.exists()


def test_safe_rename_never_overwrites_on_collision(tmp_path):
    existing = _make_docx(tmp_path / "Q3 Marketing Budget Proposal.docx", paragraphs=["original"])
    path = _make_docx(tmp_path / "doc2.docx", paragraphs=["different content"])
    log_path = tmp_path / "log.json"

    result = safe_rename(path, "Q3 Marketing Budget Proposal", log_path)

    assert result.status == "renamed"
    assert result.new_path.name == "Q3 Marketing Budget Proposal-2.docx"
    # the original file with that name is untouched
    assert existing.exists()
    assert "original" in extract_docx(existing).body_text


def test_safe_rename_logs_every_rename(tmp_path):
    path = _make_docx(tmp_path / "doc1.docx", paragraphs=["content"])
    log_path = tmp_path / "log.json"
    safe_rename(path, "New Name", log_path)

    import json
    entries = json.loads(log_path.read_text(encoding="utf-8"))
    assert len(entries) == 1
    assert entries[0]["old_path"].endswith("doc1.docx")
    assert entries[0]["new_path"].endswith("New Name.docx")
