import uuid
from pathlib import Path

from docx import Document

from app.database import init_db, create_task, get_task, get_logs, get_actions
import app.organizer.task_runner as task_runner
from app.organizer.task_runner import (
    resolve_folder,
    run_organizer_task,
    FolderResolutionError,
    _extract_explicit_docx_target,
    _extract_explicit_filename,
)


def _new_task_id() -> str:
    # This suite runs against the real dev DB (agent.db), which persists
    # across test runs -- a fixed task_id would collide with a row inserted
    # by a previous run (sqlite3.IntegrityError: UNIQUE constraint failed).
    return f"task_test_organizer_{uuid.uuid4().hex[:8]}"


def _make_docx(path: Path, title: str = "", paragraphs=None):
    doc = Document()
    if title:
        doc.core_properties.title = title
    for p in paragraphs or []:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def setup_module(module):
    # Tables are CREATE TABLE IF NOT EXISTS -- safe to call against the real
    # dev DB without disturbing existing rows.
    init_db()


# --- resolve_folder --------------------------------------------------------

def test_resolve_folder_finds_explicit_absolute_path(tmp_path):
    assert resolve_folder(f"rename the files in {tmp_path}") == tmp_path


def test_resolve_folder_raises_on_nonexistent_explicit_path():
    try:
        resolve_folder("rename the files in /this/path/does/not/exist/anywhere")
        assert False, "expected FolderResolutionError"
    except FolderResolutionError:
        pass


def test_resolve_folder_raises_when_nothing_identifiable():
    try:
        resolve_folder("please rename my files")
        assert False, "expected FolderResolutionError"
    except FolderResolutionError:
        pass


def test_resolve_folder_matches_known_alias_when_present(tmp_path, monkeypatch):
    # Point resolution at a fake "home" so this doesn't depend on the actual
    # environment having a Downloads folder.
    downloads = tmp_path / "Downloads"
    downloads.mkdir()
    monkeypatch.setattr(task_runner, "_onedrive_roots", lambda: [tmp_path])
    assert resolve_folder("rename the doc1 files in my Downloads folder") == downloads


def test_resolve_folder_finds_onedrive_redirected_desktop(tmp_path, monkeypatch):
    """Regression test for a real bug: OneDrive's "Backup"/"Known Folder Move"
    feature redirects Desktop (and Documents/Pictures) to live under the
    OneDrive folder instead of directly under the user's home directory, so
    Path.home() / "Desktop" doesn't exist even on a machine with a real,
    populated Desktop -- it's just not where Windows normally puts it. This
    is exactly what happened when a user ran 'On the desktop we have Guide
    Acceptance word doc... rename the file accordingly' and got 'Could not
    find your Desktop folder at the expected location (C:\\Users\\<user>\\Desktop)'
    despite genuinely having a Desktop folder, just under OneDrive."""
    home = tmp_path / "home"
    home.mkdir()
    onedrive_desktop = home / "OneDrive" / "Desktop"
    onedrive_desktop.mkdir(parents=True)
    # plain home/Desktop deliberately does NOT exist -- that's the bug scenario

    monkeypatch.setattr(task_runner.Path, "home", staticmethod(lambda: home))
    assert resolve_folder("rename the file on my desktop based on its content") == onedrive_desktop


def test_resolve_folder_error_lists_every_location_checked(tmp_path, monkeypatch):
    home = tmp_path / "home"
    home.mkdir()
    # no Desktop anywhere -- plain home or under OneDrive
    (home / "OneDrive").mkdir()
    monkeypatch.setattr(task_runner.Path, "home", staticmethod(lambda: home))

    try:
        resolve_folder("rename the file on my desktop")
        assert False, "expected FolderResolutionError"
    except FolderResolutionError as e:
        assert str(home / "Desktop") in str(e)
        assert str(home / "OneDrive" / "Desktop") in str(e)


# --- explicit filename extraction -------------------------------------------

def test_extract_explicit_filename_finds_bare_token_in_sentence():
    # This is the exact real-world prompt that exposed the bug: the token
    # "bul.docx" appears mid-sentence, surrounded by ordinary words, and the
    # word "docx" (no dot) appears again later without a filename attached.
    prompt = "on desktop you can find bul.docx rename it based on the content inside the docx"
    assert _extract_explicit_filename(prompt) == "bul.docx"


def test_extract_explicit_filename_returns_none_when_absent():
    assert _extract_explicit_filename("rename the doc1 files in Downloads") is None


def test_extract_explicit_docx_target_finds_absolute_path(tmp_path):
    prompt = f"please rename {tmp_path / 'bul.docx'} based on its content"
    assert _extract_explicit_docx_target(prompt) == tmp_path / "bul.docx"


def test_extract_explicit_docx_target_ignores_folder_only_paths(tmp_path):
    prompt = f"rename the files in {tmp_path}"
    assert _extract_explicit_docx_target(prompt) is None


# --- run_organizer_task ------------------------------------------------------

def test_run_organizer_task_renames_explicitly_named_file_even_if_not_generic(tmp_path):
    """Regression test for the real bug report: a user had a file named
    'bul.docx' (not one of the recognized generic patterns like doc1.docx)
    on their Desktop and asked the agent to 'rename bul.docx based on the
    content inside the docx'. The folder-scan path silently skipped it
    because is_generically_named('bul') is False -- but naming the file
    directly should be enough authorization on its own, regardless of what
    it's currently called."""
    _make_docx(tmp_path / "bul.docx", title="Bulletin Board Meeting Notes", paragraphs=["content"])

    task_id = _new_task_id()
    prompt = f"on desktop you can find bul.docx rename it based on the content inside the docx, folder is {tmp_path}"
    create_task(task_id, prompt)
    run_organizer_task(task_id, prompt)

    task = get_task(task_id)
    assert task["status"] == "completed"
    assert not (tmp_path / "bul.docx").exists()
    assert (tmp_path / "Bulletin Board Meeting Notes.docx").exists()


def test_run_organizer_task_renames_via_absolute_file_path(tmp_path):
    target = _make_docx(tmp_path / "bul.docx", title="Bulletin Board Meeting Notes", paragraphs=["content"])

    task_id = _new_task_id()
    prompt = f"please rename {target} based on its content"
    create_task(task_id, prompt)
    run_organizer_task(task_id, prompt)

    task = get_task(task_id)
    assert task["status"] == "completed"
    assert (tmp_path / "Bulletin Board Meeting Notes.docx").exists()


def test_run_organizer_task_fails_clearly_when_named_file_not_found(tmp_path):
    task_id = _new_task_id()
    prompt = f"rename bul.docx based on its content, folder is {tmp_path}"
    create_task(task_id, prompt)
    run_organizer_task(task_id, prompt)

    task = get_task(task_id)
    assert task["status"] == "failed"
    assert "bul.docx" in task["error"]


def test_run_organizer_task_still_only_renames_generic_files_without_explicit_name(tmp_path):
    """Make sure the fix for explicitly-named files didn't loosen the
    folder-scan path -- a prompt with no filename mentioned should still only
    touch generically-named files, same as before."""
    _make_docx(tmp_path / "Family Vacation Itinerary.docx", paragraphs=["not touched"])

    task_id = _new_task_id()
    prompt = f"rename the files in {tmp_path}"
    create_task(task_id, prompt)
    run_organizer_task(task_id, prompt)

    task = get_task(task_id)
    assert task["status"] == "completed"
    assert (tmp_path / "Family Vacation Itinerary.docx").exists()


def test_run_organizer_task_renames_table_only_document_based_on_content(tmp_path):
    """End-to-end regression test for the table-content extraction fix: a
    document whose only content lives in a table (no free-standing
    paragraphs) should be named from that content, not fall back to a
    timestamped placeholder."""
    from docx import Document

    path = tmp_path / "bul.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Quarterly Budget Review"
    table.cell(0, 1).text = "approved by finance committee"
    doc.save(str(path))

    task_id = _new_task_id()
    prompt = f"rename bul.docx based on its content, folder is {tmp_path}"
    create_task(task_id, prompt)
    run_organizer_task(task_id, prompt)

    task = get_task(task_id)
    assert task["status"] == "completed"
    assert not path.exists()
    remaining = list(tmp_path.glob("*.docx"))
    assert len(remaining) == 1
    assert not remaining[0].stem.startswith("Untitled Document")


def test_run_organizer_task_logs_diagnostic_when_no_content_to_name_from(tmp_path):
    """When a document genuinely has nothing to name itself from (no title,
    no headings, no paragraphs, no tables), the task should still complete
    with a placeholder name, but the log should explain why instead of
    leaving the user to guess."""
    from docx import Document

    path = tmp_path / "bul.docx"
    Document().save(str(path))  # a totally blank document

    task_id = _new_task_id()
    prompt = f"rename bul.docx based on its content, folder is {tmp_path}"
    create_task(task_id, prompt)
    run_organizer_task(task_id, prompt)

    task = get_task(task_id)
    assert task["status"] == "completed"
    remaining = list(tmp_path.glob("*.docx"))
    assert len(remaining) == 1
    assert remaining[0].stem.startswith("Untitled Document")

    logs = get_logs(task_id)
    assert any("Could not find a title, heading, or usable content" in l["message"] for l in logs)


def test_run_organizer_task_renames_generic_files_and_completes(tmp_path):
    _make_docx(tmp_path / "doc1.docx", title="Q3 Budget Review", paragraphs=["content"])
    _make_docx(tmp_path / "Family Vacation Itinerary.docx", paragraphs=["not touched"])

    task_id = _new_task_id()
    create_task(task_id, f"rename the files in {tmp_path}")
    run_organizer_task(task_id, f"rename the files in {tmp_path}")

    task = get_task(task_id)
    assert task["status"] == "completed"

    # generic file was renamed based on its title metadata
    assert not (tmp_path / "doc1.docx").exists()
    assert (tmp_path / "Q3 Budget Review.docx").exists()
    # non-generic file was left alone
    assert (tmp_path / "Family Vacation Itinerary.docx").exists()

    actions = get_actions(task_id)
    assert any(a["action_type"] == "rename" for a in actions)


def test_run_organizer_task_completes_with_no_candidates(tmp_path):
    _make_docx(tmp_path / "Family Vacation Itinerary.docx", paragraphs=["not touched"])

    task_id = _new_task_id()
    create_task(task_id, f"rename the files in {tmp_path}")
    run_organizer_task(task_id, f"rename the files in {tmp_path}")

    task = get_task(task_id)
    assert task["status"] == "completed"
    assert (tmp_path / "Family Vacation Itinerary.docx").exists()


def test_run_organizer_task_fails_clearly_when_folder_not_identifiable():
    task_id = _new_task_id()
    create_task(task_id, "please rename my files")
    run_organizer_task(task_id, "please rename my files")

    task = get_task(task_id)
    assert task["status"] == "failed"
    assert task["error"]

    logs = get_logs(task_id)
    assert any("couldn't tell which folder" in l["message"] for l in logs)
